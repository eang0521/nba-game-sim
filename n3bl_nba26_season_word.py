import json
import random
import os
import sys
from collections import defaultdict

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("python-docx is not installed. Run: pip install python-docx")
    sys.exit(1)

def roll_d(sides, num=1):
    return sum(random.randint(1, sides) for _ in range(num))

# --- LEAGUE STRUCTURE ---
DIVISIONS = {
    "BOS": "Atlantic", "BKN": "Atlantic", "NYK": "Atlantic", "PHI": "Atlantic", "TOR": "Atlantic",
    "CHI": "Central", "CLE": "Central", "DET": "Central", "IND": "Central", "MIL": "Central",
    "ATL": "Southeast", "CHA": "Southeast", "MIA": "Southeast", "ORL": "Southeast", "WAS": "Southeast",
    "DEN": "Northwest", "MIN": "Northwest", "OKC": "Northwest", "POR": "Northwest", "UTA": "Northwest",
    "DAL": "Southwest", "HOU": "Southwest", "MEM": "Southwest", "NOP": "Southwest", "SAS": "Southwest",
    "GSW": "Pacific", "LAC": "Pacific", "LAL": "Pacific", "PHX": "Pacific", "SAC": "Pacific"
}

CONFERENCES = {
    "Atlantic": "East", "Central": "East", "Southeast": "East",
    "Northwest": "West", "Southwest": "West", "Pacific": "West"
}

class Player:
    def __init__(self, name, t_2pt, t_3pt, t_def, t_reb):
        self.name = name
        self.t_2pt = t_2pt
        self.t_3pt = t_3pt
        self.t_def = t_def
        self.t_reb = t_reb

        self.ovr = t_2pt + t_3pt + t_def + t_reb
        self.pv = t_reb - t_3pt + (0.01 * t_2pt) + (0.0001 * t_def)
        self.off = t_2pt + t_3pt + (0.01 * self.ovr)

        self.ft_chance = max(1, min(90, round(50 + (t_2pt * 0.4) + (t_3pt * 1.0))))

        # Season Stats Tracker (Accumulates)
        self.games_played = 0
        self.s_seconds = 0
        self.s_pts = 0
        self.s_fgm = 0
        self.s_fga = 0
        self.s_tptm = 0
        self.s_tpta = 0
        self.s_ftm = 0
        self.s_fta = 0
        self.s_reb = 0
        self.s_ast = 0
        self.s_stl = 0
        self.s_blk = 0
        self.s_pf = 0
        self.s_pm = 0

        self.reset_game_state()

    def reset_game_state(self):
        # Game Stats Tracker (Resets every game)
        self.seconds_played = 0
        self.points = 0
        self.fg_made = 0
        self.fg_attempted = 0
        self.tpt_made = 0
        self.tpt_attempted = 0
        self.ft_made = 0
        self.ft_attempted = 0
        self.rebounds = 0
        self.assists = 0
        self.steals = 0
        self.blocks = 0
        self.fouls = 0
        self.plus_minus = 0

        # In-Game State
        self.on_court = False
        self.fouled_out = False
        self.positions_played = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0}
        self.current_position = None

    _SEASON_GAME_PAIRS = [
        ("s_seconds", "seconds_played"), ("s_pts", "points"),
        ("s_fgm", "fg_made"), ("s_fga", "fg_attempted"),
        ("s_tptm", "tpt_made"), ("s_tpta", "tpt_attempted"),
        ("s_ftm", "ft_made"), ("s_fta", "ft_attempted"),
        ("s_reb", "rebounds"), ("s_ast", "assists"),
        ("s_stl", "steals"), ("s_blk", "blocks"),
        ("s_pf", "fouls"), ("s_pm", "plus_minus"),
    ]

    def record_game(self):
        if sum(self.positions_played.values()) > 0:
            self.games_played += 1
            for season_attr, game_attr in self._SEASON_GAME_PAIRS:
                setattr(self, season_attr, getattr(self, season_attr) + getattr(self, game_attr))

    def get_primary_position_name(self):
        if sum(self.positions_played.values()) == 0:
            return "BN"
        best_pos = max(self.positions_played, key=self.positions_played.get)
        return {1: "PG", 2: "SG", 3: "SF", 4: "PF", 5: "C"}[best_pos]

class Team:
    def __init__(self, abbr, player_data):
        self.abbr = abbr
        self.players = [Player(p['name'], p['t_2pt'], p['t_3pt'], p['t_def'], p['t_reb']) for p in player_data]
        self.division = DIVISIONS.get(abbr, "Unknown")
        self.conference = CONFERENCES.get(self.division, "Unknown")

        # Season Stats
        self.wins = 0
        self.losses = 0
        self.points_scored = 0
        self.points_allowed = 0
        self.h2h_wins = defaultdict(int)
        self.h2h_games = defaultdict(int)
        self.div_wins = 0
        self.div_losses = 0
        self.conf_wins = 0
        self.conf_losses = 0

        # In-Game State
        self.score = 0
        self.court = []
        self.defensive_fouls_qtr = 0
        self.foul_in_last_2_mins = False

    def reset_game_state(self):
        self.score = 0
        self.court = []
        self.defensive_fouls_qtr = 0
        self.foul_in_last_2_mins = False
        for p in self.players:
            p.reset_game_state()

    def update_season_stats(self, opponent, won, pts_scored, pts_allowed):
        if won:
            self.wins += 1
            self.h2h_wins[opponent.abbr] += 1
            if self.division == opponent.division: self.div_wins += 1
            if self.conference == opponent.conference: self.conf_wins += 1
        else:
            self.losses += 1
            if self.division == opponent.division: self.div_losses += 1
            if self.conference == opponent.conference: self.conf_losses += 1

        self.h2h_games[opponent.abbr] += 1
        self.points_scored += pts_scored
        self.points_allowed += pts_allowed

        for p in self.players:
            p.record_game()

    @property
    def win_pct(self):
        return self.wins / (self.wins + self.losses) if (self.wins + self.losses) > 0 else 0

    def get_court_sorted_by_pv(self):
        sorted_court = sorted(self.court, key=lambda x: x.pv)
        for i, p in enumerate(sorted_court):
            p.current_position = i + 1
        return sorted_court

    def get_court_sorted_by_off(self):
        return sorted(self.court, key=lambda x: x.off, reverse=True)


class Game:
    def __init__(self, team1, team2, verbose=False):
        self.t1 = team1
        self.t2 = team2
        self.verbose = verbose
        self.quarter = 0
        self.clock = 0
        self.possession = None
        self.jumpball_winner = None
        self.subs_waiting = False

        self.t1.reset_game_state()
        self.t2.reset_game_state()

        self.t1_q_scores = []
        self.t2_q_scores = []
        self.t1_last_q_score = 0
        self.t2_last_q_score = 0

    def format_time(self):
        mins = self.clock // 60
        secs = self.clock % 60
        return f"{mins:02}:{secs:02}"

    def log(self, message):
        if not self.verbose: return
        score_str = f"[{self.t1.abbr} {self.t1.score} - {self.t2.abbr} {self.t2.score}]" if "makes" in message or "scores" in message else ""
        print(f"{self.format_time()} {message} {score_str}".strip())

    def update_plus_minus(self, scoring_team, points):
        other_team = self.t2 if scoring_team == self.t1 else self.t1
        for p in scoring_team.court:
            p.plus_minus += points
        for p in other_team.court:
            p.plus_minus -= points

    def start_quarter_lineups(self):
        for team in [self.t1, self.t2]:
            team.court = []
            for p in team.players:
                p.on_court = False

            for _ in range(5):
                bench = sorted([p for p in team.players if not p.on_court and not p.fouled_out], key=lambda x: x.ovr, reverse=True)
                if not bench: break

                selected = bench[self.get_threshold_index(roll_d(20), len(bench))]
                selected.on_court = True
                team.court.append(selected)

            team.get_court_sorted_by_pv()

    def display_lineups(self):
        if not self.verbose: return
        t1_sorted = self.t1.get_court_sorted_by_pv()
        t2_sorted = self.t2.get_court_sorted_by_pv()

        print("\nCURRENT LINEUPS:")
        print(f"{self.t1.abbr:<28} | {self.t2.abbr:<28}")
        print("-" * 59)
        positions = ["PG", "SG", "SF", "PF", "C"]

        max_court = max(len(t1_sorted), len(t2_sorted))
        for i in range(max_court):
            pos = positions[i] if i < 5 else "Ext"
            p1_str = f"{pos}: {t1_sorted[i].name}" if i < len(t1_sorted) else "NONE"
            p2_str = f"{pos}: {t2_sorted[i].name}" if i < len(t2_sorted) else "NONE"
            print(f"{p1_str:<28} | {p2_str:<28}")
        print()

    def get_threshold_index(self, roll, pool_size):
        if pool_size >= 5: thresholds = [6, 11, 15, 18, 20]
        elif pool_size == 4: thresholds = [8, 14, 18, 20]
        elif pool_size == 3: thresholds = [9, 16, 20]
        elif pool_size == 2: thresholds = [13, 20]
        else: thresholds = [20]

        for i, t in enumerate(thresholds):
            if roll <= t: return i
        return len(thresholds) - 1

    def handle_substitutions(self, ft_shooter=None, forced_sub_players=None):
        if forced_sub_players is None: forced_sub_players = []
        subs_made = False

        for team in [self.t1, self.t2]:
            forced_team = [p for p in forced_sub_players if p in team.players and p.on_court]
            if not self.subs_waiting and not forced_team: continue

            num_subs = 1 if forced_team and not self.subs_waiting else min(5, self.get_threshold_index(roll_d(20), 5) + 1)
            num_subs = max(num_subs, len(forced_team))

            court_pool = sorted(team.court, key=lambda x: x.ovr)
            available_bench_total = [p for p in team.players if not p.on_court and not p.fouled_out]

            out_players, in_players = [], []

            for i in range(min(num_subs, len(court_pool), len(available_bench_total))):
                if i < len(forced_team): out_p = forced_team[i]
                else:
                    valid_out = [p for p in court_pool if p not in out_players and p != ft_shooter]
                    if not valid_out: break
                    out_p = valid_out[self.get_threshold_index(roll_d(20), len(valid_out))]

                valid_in_all = sorted([p for p in team.players if not p.on_court and not p.fouled_out and p not in in_players], key=lambda x: x.ovr, reverse=True)
                valid_in = valid_in_all[:5]
                if not valid_in: break
                in_p = valid_in[min(self.get_threshold_index(roll_d(20), 5), len(valid_in) - 1)]

                out_players.append(out_p)
                in_players.append(in_p)

            for op, ip in zip(out_players, in_players):
                op.on_court = False
                team.court.remove(op)
                ip.on_court = True
                team.court.append(ip)
                if self.verbose: self.log(f"SUB: {ip.name} in for {op.name}")
                subs_made = True

        self.subs_waiting = False
        self.t1.get_court_sorted_by_pv()
        self.t2.get_court_sorted_by_pv()
        if subs_made: self.display_lineups()

    def advance_time(self, dice_count):
        time_taken = roll_d(6, dice_count)
        self.clock -= time_taken
        if self.clock < 0: self.clock = 0

        for team in [self.t1, self.t2]:
            for p in team.court:
                p.seconds_played += time_taken
                p.positions_played[p.current_position] += time_taken
        return self.clock > 0

    def record_foul(self, player, defense_team):
        player.fouls += 1
        forced_sub = []
        if player.fouls >= 6:
            player.fouled_out = True
            forced_sub.append(player)
            self.log(f"{player.name} fouled out!")
        elif player.fouls >= (self.quarter + 2):
            forced_sub.append(player)
            self.log(f"{player.name} is in foul trouble ({player.fouls} fouls) and must sub out.")
        return forced_sub

    def shoot_fts(self, shooter, num_fts, team):
        if num_fts == 0: return True
        made_last = False
        for i in range(num_fts):
            shooter.ft_attempted += 1
            if roll_d(100) <= shooter.ft_chance:
                shooter.points += 1
                shooter.ft_made += 1
                team.score += 1
                self.update_plus_minus(team, 1)
                self.log(f"{shooter.name} makes free throw {i+1}/{num_fts}")
                if i == num_fts - 1: made_last = True
            else:
                self.log(f"{shooter.name} misses free throw {i+1}/{num_fts}")
                if i == num_fts - 1: made_last = False
        return made_last

    def play_possession(self, dice_count=4):
        if self.clock <= 0: return
        if not self.advance_time(dice_count): return

        off_team = self.possession
        def_team = self.t2 if off_team == self.t1 else self.t1

        event = roll_d(6)
        off_court_pv = off_team.get_court_sorted_by_pv()
        def_court_pv = def_team.get_court_sorted_by_pv()

        if not off_court_pv or not def_court_pv:
            return

        if event == 1:
            for defender in def_court_pv:
                safe_def = max(1, defender.t_def)
                if roll_d(100) <= safe_def:
                    defender.steals += 1
                    self.log(f"{defender.name} Steal")
                    self.possession = def_team
                    return
            event = random.randint(3, 6)

        if event == 2:
            foul_type = roll_d(10)
            foul_pos = min((roll_d(10) - 1) // 2, len(off_court_pv) - 1)

            if foul_type == 1: # Offensive
                fouler = off_court_pv[foul_pos]
                forced = self.record_foul(fouler, def_team)
                self.log(f"{fouler.name} offensive foul")
                bonus = (def_team.defensive_fouls_qtr >= 4) or (self.clock <= 120 and off_team.foul_in_last_2_mins)
                if bonus:
                    shooter = random.choice(def_team.court)
                    self.handle_substitutions(ft_shooter=shooter, forced_sub_players=forced)
                    if not self.shoot_fts(shooter, 2, def_team): self.do_rebound_cycle(def_team, off_team, skipped_shooter=shooter)
                    else: self.possession = off_team
                else:
                    self.handle_substitutions(forced_sub_players=forced)
                    self.possession = def_team
            else: # Defensive
                foul_pos = min(foul_pos, len(def_court_pv) - 1)
                fouler = def_court_pv[foul_pos]
                forced = self.record_foul(fouler, def_team)
                def_team.defensive_fouls_qtr += 1
                if self.clock <= 120: def_team.foul_in_last_2_mins = True

                if foul_type <= 6: # On floor
                    self.log(f"{fouler.name} defensive foul")
                    bonus = (def_team.defensive_fouls_qtr >= 4) or (self.clock <= 120 and def_team.foul_in_last_2_mins)
                    if bonus:
                        shooter = random.choice(off_team.court)
                        self.handle_substitutions(ft_shooter=shooter, forced_sub_players=forced)
                        if not self.shoot_fts(shooter, 2, off_team): self.do_rebound_cycle(off_team, def_team, skipped_shooter=shooter)
                        else: self.possession = def_team
                    else:
                        self.handle_substitutions(forced_sub_players=forced)
                        self.play_possession(dice_count=2)
                else: # Shooting
                    self.log(f"{fouler.name} shooting foul")
                    off_court_off = off_team.get_court_sorted_by_off()
                    shooter = off_court_off[self.get_threshold_index(roll_d(20), len(off_court_off))]

                    safe_2pt, safe_3pt = max(1, shooter.t_2pt), max(1, shooter.t_3pt)
                    shot_ratio = round((safe_2pt * 3) / ((safe_2pt * 3) + (safe_3pt * 2)) * 100)
                    is_2pt = roll_d(100) <= shot_ratio
                    if not is_2pt and roll_d(6) <= 4: is_2pt = True

                    pts = 2 if is_2pt else 3
                    defender = def_court_pv[min(shooter.current_position - 1, len(def_court_pv) - 1)]
                    diff = (shooter.t_2pt if is_2pt else shooter.t_3pt) - defender.t_def
                    target = round(diff * ((2/3) if is_2pt else (1/2))) + (54 if is_2pt else 36)

                    if roll_d(100) <= target and roll_d(100) <= target: # And-1
                        shooter.fg_attempted += 1
                        shooter.fg_made += 1
                        if not is_2pt:
                            shooter.tpt_attempted += 1
                            shooter.tpt_made += 1
                        shooter.points += pts
                        off_team.score += pts
                        self.update_plus_minus(off_team, pts)

                        assister = self._pick_assister(shooter, off_court_pv)
                        self.log(f"{shooter.name} makes {pts} (And-1!)" + (f" ({assister.name} Ast)" if assister else ""))
                        if roll_d(10) <= pts: self.subs_waiting = True
                        self.handle_substitutions(ft_shooter=shooter, forced_sub_players=forced)
                        if not self.shoot_fts(shooter, 1, off_team): self.do_rebound_cycle(off_team, def_team, skipped_shooter=shooter)
                        else: self.possession = def_team
                    else: # Missed
                        self.log(f"{shooter.name} misses {pts} on foul")
                        self.handle_substitutions(ft_shooter=shooter, forced_sub_players=forced)
                        if not self.shoot_fts(shooter, pts, off_team): self.do_rebound_cycle(off_team, def_team, skipped_shooter=shooter)
                        else: self.possession = def_team
            return

        if event >= 3:
            off_court_off = off_team.get_court_sorted_by_off()
            self._resolve_shot(off_court_off, off_court_pv, def_court_pv, off_team, def_team)

    def do_rebound_cycle(self, off_team, def_team, skipped_shooter):
        off_pv = off_team.get_court_sorted_by_pv()
        def_pv = def_team.get_court_sorted_by_pv()

        cycle_num = 1
        while self.clock > 0:
            for reb_player in list(reversed(def_pv)) + list(reversed(off_pv)):
                if cycle_num == 1 and reb_player == skipped_shooter: continue
                reb_rating = max(1, reb_player.t_reb) + (3 if cycle_num == 1 and reb_player in def_pv else 0)
                if roll_d(100) <= reb_rating:
                    reb_player.rebounds += 1
                    self.log(f"{reb_player.name} Rebound")
                    if reb_player in def_pv: self.possession = def_team
                    else:
                        self.possession = off_team
                        self.play_possession(dice_count=2)
                    return
            self.clock -= 1
            cycle_num += 1

    def _pick_assister(self, shooter, off_court_pv):
        assist_roll = roll_d(100)
        ast_idx = 0 if assist_roll <= 24 else 1 if assist_roll <= 42 else 2 if assist_roll <= 54 else 3 if assist_roll <= 60 else -1
        if ast_idx == -1:
            return None
        tmts = [p for p in off_court_pv if p != shooter]
        if ast_idx < len(tmts):
            assister = tmts[ast_idx]
            assister.assists += 1
            return assister
        return None

    def _resolve_shot(self, off_court_off, off_court_pv, def_court_pv, off_team, def_team):
        shooter = off_court_off[self.get_threshold_index(roll_d(20), len(off_court_off))]
        safe_2pt, safe_3pt = max(1, shooter.t_2pt), max(1, shooter.t_3pt)
        shot_ratio = round((safe_2pt * 3) / ((safe_2pt * 3) + (safe_3pt * 2)) * 100)
        is_2pt = roll_d(100) <= shot_ratio
        pts = 2 if is_2pt else 3
        defender = def_court_pv[min(shooter.current_position - 1, len(def_court_pv) - 1)]
        diff = (shooter.t_2pt if is_2pt else shooter.t_3pt) - defender.t_def
        target = round(diff * ((2/3) if is_2pt else (1/2))) + (54 if is_2pt else 36)

        shooter.fg_attempted += 1
        if not is_2pt:
            shooter.tpt_attempted += 1

        if roll_d(100) <= target:
            shooter.fg_made += 1
            if not is_2pt:
                shooter.tpt_made += 1
            shooter.points += pts
            off_team.score += pts
            self.update_plus_minus(off_team, pts)

            assister = self._pick_assister(shooter, off_court_pv)
            self.log(f"{shooter.name} makes {pts}" + (f" ({assister.name} Ast)" if assister else ""))
            if roll_d(10) <= pts:
                self.subs_waiting = True
            self.possession = def_team
        else:
            blocked = False
            if roll_d(100) <= max(1, defender.t_def):
                if roll_d(6) <= defender.current_position + 1:
                    blocked = True
                    defender.blocks += 1
                    self.log(f"{shooter.name} misses {pts} (Blocked by {defender.name}!)")
            if not blocked:
                self.log(f"{shooter.name} misses {pts}")
            self.do_rebound_cycle(off_team, def_team, skipped_shooter=shooter)

    def play_quarter(self, mins):
        self.subs_waiting = False
        self.clock = mins * 60
        self.t1.defensive_fouls_qtr = self.t2.defensive_fouls_qtr = 0
        self.t1.foul_in_last_2_mins = self.t2.foul_in_last_2_mins = False

        self.start_quarter_lineups()
        self.display_lineups()

        if self.quarter == 1 or mins == 5:
            c1 = self.t1.get_court_sorted_by_pv()[-1] if self.t1.court else None
            c2 = self.t2.get_court_sorted_by_pv()[-1] if self.t2.court else None

            if c1 and c2:
                if roll_d(max(1, c1.t_reb) + max(1, c2.t_reb)) <= max(1, c1.t_reb):
                    self.possession, self.jumpball_winner = self.t1, (self.t1 if self.quarter == 1 else self.jumpball_winner)
                    self.log(f"{c1.name} wins the jumpball for {self.t1.abbr}")
                else:
                    self.possession, self.jumpball_winner = self.t2, (self.t2 if self.quarter == 1 else self.jumpball_winner)
                    self.log(f"{c2.name} wins the jumpball for {self.t2.abbr}")
            else:
                self.possession = self.t1
        elif self.quarter in [2, 3]:
            self.possession = self.t2 if self.jumpball_winner == self.t1 else self.t1
        else:
            self.possession = self.jumpball_winner

        while self.clock > 0: self.play_possession()

        # Update quarter scores
        self.t1_q_scores.append(self.t1.score - self.t1_last_q_score)
        self.t2_q_scores.append(self.t2.score - self.t2_last_q_score)
        self.t1_last_q_score = self.t1.score
        self.t2_last_q_score = self.t2.score

    def simulate(self):
        for q in range(1, 5):
            self.quarter = q
            if self.verbose: print(f"\n--- START OF QUARTER {q} ---")
            self.play_quarter(12)

        ot = 1
        while self.t1.score == self.t2.score:
            if self.verbose: print(f"\n--- START OF OVERTIME {ot} ---")
            self.quarter = 4 + ot
            self.play_quarter(5)
            ot += 1

        t1_won = self.t1.score > self.t2.score
        self.t1.update_season_stats(self.t2, t1_won, self.t1.score, self.t2.score)
        self.t2.update_season_stats(self.t1, not t1_won, self.t2.score, self.t1.score)

        if self.verbose:
            print(f"\nFINAL SCORE: {self.t1.abbr} {self.t1.score} - {self.t2.abbr} {self.t2.score}")
        return self.t1 if t1_won else self.t2

    def format_box_score(self):
        lines = []
        lines.append(f"FINAL SCORE: {self.t1.abbr} {self.t1.score} - {self.t2.abbr} {self.t2.score}")

        headers = ["Team"] + [f"Q{i+1}" for i in range(4)]
        if len(self.t1_q_scores) > 4:
            headers += [f"OT{i-3}" for i in range(4, len(self.t1_q_scores))]
        headers.append("T")

        lines.append("")
        lines.append(f"{headers[0]:<5}| " + " | ".join(f"{h:>3}" for h in headers[1:]))

        def linescore(team, scores):
            return f"{team.abbr:<5}| " + " | ".join(f"{s:>3}" for s in scores) + f" | {team.score:>3}"

        lines.append(linescore(self.t1, self.t1_q_scores))
        lines.append(linescore(self.t2, self.t2_q_scores))
        lines.append("")
        lines.append(f"{'Player':<28} | {'MIN':>3} | {'PTS':>3} | {'FG':>5} | {'3PT':>5} | {'FT':>5} | {'REB':>3} | {'AST':>3} | {'STL':>3} | {'BLK':>3} | {'PF':>2} | {'+/-':>4} |")
        lines.append("-" * 102)

        for team in [self.t1, self.t2]:
            lines.append(f"--- {team.abbr} ---")
            played = sorted([p for p in team.players if p.seconds_played > 0],
                            key=lambda x: x.seconds_played, reverse=True)
            for p in played:
                mp = round(p.seconds_played / 60)
                pm = f"+{p.plus_minus}" if p.plus_minus > 0 else str(p.plus_minus)
                name_pos = f"{p.name} ({p.get_primary_position_name()})"
                lines.append(f"{name_pos:<28} | {mp:>3} | {p.points:>3} | {f'{p.fg_made}-{p.fg_attempted}':>5} | {f'{p.tpt_made}-{p.tpt_attempted}':>5} | {f'{p.ft_made}-{p.ft_attempted}':>5} | {p.rebounds:>3} | {p.assists:>3} | {p.steals:>3} | {p.blocks:>3} | {p.fouls:>2} | {pm:>4} |")
            lines.append("-" * 102)
            s = lambda attr: sum(getattr(p, attr) for p in played)
            t_min = sum(round(p.seconds_played / 60) for p in played)
            t_pm = s("plus_minus")
            t_pm_str = f"+{t_pm}" if t_pm > 0 else str(t_pm)
            t_fg = f"{s('fg_made')}-{s('fg_attempted')}"
            t_tpt = f"{s('tpt_made')}-{s('tpt_attempted')}"
            t_ft = f"{s('ft_made')}-{s('ft_attempted')}"
            lines.append(f"{'TOTALS':<28} | {t_min:>3} | {s('points'):>3} | {t_fg:>5} | {t_tpt:>5} | {t_ft:>5} | {s('rebounds'):>3} | {s('assists'):>3} | {s('steals'):>3} | {s('blocks'):>3} | {s('fouls'):>2} | {t_pm_str:>4} |")
            lines.append("")
        return lines

    def print_box_score(self):
        for line in self.format_box_score():
            print(line)


class SeasonMode:
    def __init__(self, teams_data):
        self.teams = {abbr: Team(abbr, data) for abbr, data in teams_data.items()}
        self.schedule = []
        self.box_scores = []

    def build_schedule(self):
        """Builds exactly 82 games per team according to NBA rules."""
        games_needed = defaultdict(lambda: defaultdict(int))
        for t1 in self.teams.values():
            for t2 in self.teams.values():
                if t1 == t2: continue
                if t1.conference != t2.conference:
                    games_needed[t1.abbr][t2.abbr] = 2
                elif t1.division == t2.division:
                    games_needed[t1.abbr][t2.abbr] = 4
                else:
                    games_needed[t1.abbr][t2.abbr] = 3

        for conf in ["East", "West"]:
            divs = list(set(t.division for t in self.teams.values() if t.conference == conf))
            div_pairings = [(divs[0], divs[1]), (divs[1], divs[2]), (divs[2], divs[0])]

            for d1, d2 in div_pairings:
                t1s = [t for t in self.teams.values() if t.division == d1]
                t2s = [t for t in self.teams.values() if t.division == d2]

                for i in range(5):
                    for j in range(3):
                        opp = t2s[(i + j) % 5]
                        games_needed[t1s[i].abbr][opp.abbr] = 4
                        games_needed[opp.abbr][t1s[i].abbr] = 4

        for t1, opponents in games_needed.items():
            for t2 in list(opponents.keys()):
                while games_needed[t1][t2] > 0:
                    self.schedule.append((t1, t2))
                    games_needed[t1][t2] -= 1
                    games_needed[t2][t1] -= 1

    def sim_regular_season(self):
        print("Simulating 82-Game Regular Season. Please wait...")
        random.shuffle(self.schedule)
        for i, (a_abbr, b_abbr) in enumerate(self.schedule, 1):
            sys.stderr.write(f"\rGame {i}/{len(self.schedule)}...")
            sys.stderr.flush()
            game = Game(self.teams[a_abbr], self.teams[b_abbr], verbose=False)
            game.simulate()
            self.box_scores.append(f"{'='*60}")
            self.box_scores.append(f"GAME {i}: {a_abbr} vs {b_abbr}")
            self.box_scores.append(f"{'='*60}")
            self.box_scores.extend(game.format_box_score())
        sys.stderr.write("\n")

    def rank_teams(self, teams_list):
        """Tiebreaker ranking."""
        def get_tiebreaker_key(team):
            div_winner = 1 if max([t.win_pct for t in self.teams.values() if t.division == team.division]) == team.win_pct else 0
            div_pct = team.div_wins / max(1, team.div_wins + team.div_losses)
            conf_pct = team.conf_wins / max(1, team.conf_wins + team.conf_losses)
            pt_diff = team.points_scored - team.points_allowed
            return (team.win_pct, div_winner, div_pct, conf_pct, pt_diff)

        return sorted(teams_list, key=get_tiebreaker_key, reverse=True)

    def display_standings(self):
        for conf in ["East", "West"]:
            print(f"\n{conf}ern Conference Standings")
            print(f"{'Team':<5} | {'W':>2} | {'L':>2} | {'PPG':>5} | {'OPPG':>5}")
            print("-" * 33)
            conf_teams = [t for t in self.teams.values() if t.conference == conf]
            ranked = self.rank_teams(conf_teams)
            for t in ranked:
                ppg = t.points_scored / 82
                oppg = t.points_allowed / 82
                print(f"{t.abbr:<5} | {t.wins:>2} | {t.losses:>2} | {ppg:>5.1f} | {oppg:>5.1f}")

    def display_league_leaders(self):
        all_players = []
        player_team = {}
        for abbr, t in self.teams.items():
            all_players.extend(t.players)
            for p in t.players:
                player_team[p] = abbr

        def safe_div(a, b): return a / b if b > 0 else 0

        categories = {
            "MIN": lambda p: safe_div(p.s_seconds / 60, p.games_played),
            "PTS": lambda p: safe_div(p.s_pts, p.games_played),
            "REB": lambda p: safe_div(p.s_reb, p.games_played),
            "AST": lambda p: safe_div(p.s_ast, p.games_played),
            "STL": lambda p: safe_div(p.s_stl, p.games_played),
            "BLK": lambda p: safe_div(p.s_blk, p.games_played),
            "+/-": lambda p: p.s_pm,
            "FG%": lambda p: safe_div(p.s_fgm, p.s_fga) * 100,
            "3P%": lambda p: safe_div(p.s_tptm, p.s_tpta) * 100,
            "FT%": lambda p: safe_div(p.s_ftm, p.s_fta) * 100
        }

        qualifications = {
            "FG%": lambda p: p.s_fgm >= 300,
            "3P%": lambda p: p.s_tptm >= 82,
            "FT%": lambda p: p.s_ftm >= 125,
        }
        default_qual = lambda p: p.games_played >= 40

        print("\n--- LEAGUE LEADERS (Top 10) ---")
        for cat_name, key_func in categories.items():
            print(f"\n{cat_name}:")
            qual_fn = qualifications.get(cat_name, default_qual)
            qualified = [p for p in all_players if qual_fn(p)]

            leaders = sorted(qualified, key=key_func, reverse=True)[:10]
            for i, p in enumerate(leaders):
                val = key_func(p)
                fmt_val = f"{val:.1f}%" if "%" in cat_name else f"{val:.1f}" if cat_name != "+/-" else f"{val}"
                team_abbr = player_team[p]
                print(f"{i+1}. {p.name} ({team_abbr}): {fmt_val}")

    def play_series(self, t1, t2):
        t1_wins = t2_wins = 0
        game_num = 1
        while t1_wins < 4 and t2_wins < 4:
            print(f"  Game {game_num}: {t1.abbr} vs {t2.abbr}...")
            game = Game(t1, t2, verbose=False)
            winner = game.simulate()
            if winner == t1: t1_wins += 1
            else: t2_wins += 1
            self.box_scores.append(f"--- {t1.abbr} vs {t2.abbr} — Game {game_num} ---")
            self.box_scores.extend(game.format_box_score())
            game_num += 1
        series_result = f"SERIES: {t1.abbr} {t1_wins} - {t2.abbr} {t2_wins} -> {(t1 if t1_wins == 4 else t2).abbr} wins"
        print(f"  {series_result}")
        self.box_scores.append(series_result)
        self.box_scores.append("")
        return t1 if t1_wins == 4 else t2, t1_wins, t2_wins

    def run_postseason(self):
        seeds = {"East": [], "West": []}
        for conf in ["East", "West"]:
            conf_teams = [t for t in self.teams.values() if t.conference == conf]
            seeds[conf] = self.rank_teams(conf_teams)[:10]

        playoff_teams = {"East": {}, "West": {}}

        print("\n=== PLAY-IN TOURNAMENT ===")
        for conf in ["East", "West"]:
            print(f"\n{conf}ern Conference Play-In:")
            s = seeds[conf]

            # 7 vs 8
            print(f"  Play-In 7v8: {s[6].abbr} vs {s[7].abbr}...")
            game_78 = Game(s[6], s[7], verbose=False)
            w_78 = game_78.simulate()
            l_78 = s[7] if w_78 == s[6] else s[6]
            result_78 = f"7v8: {w_78.abbr} defeats {l_78.abbr} -> {w_78.abbr} clinches #7 Seed"
            print(f"  {result_78}")
            self.box_scores.append(f"--- {conf} Play-In: 7v8 — {s[6].abbr} vs {s[7].abbr} ---")
            self.box_scores.extend(game_78.format_box_score())
            self.box_scores.append(result_78)
            self.box_scores.append("")

            # 9 vs 10
            print(f"  Play-In 9v10: {s[8].abbr} vs {s[9].abbr}...")
            game_910 = Game(s[8], s[9], verbose=False)
            w_910 = game_910.simulate()
            l_910 = s[9] if w_910 == s[8] else s[8]
            result_910 = f"9v10: {w_910.abbr} defeats {l_910.abbr}"
            print(f"  {result_910}")
            self.box_scores.append(f"--- {conf} Play-In: 9v10 — {s[8].abbr} vs {s[9].abbr} ---")
            self.box_scores.extend(game_910.format_box_score())
            self.box_scores.append(result_910)
            self.box_scores.append("")

            # 8th Seed Game
            print(f"  Play-In 8th seed: {l_78.abbr} vs {w_910.abbr}...")
            game_8 = Game(l_78, w_910, verbose=False)
            w_8 = game_8.simulate()
            l_8 = w_910 if w_8 == l_78 else l_78
            result_8 = f"8th seed: {w_8.abbr} defeats {l_8.abbr} -> {w_8.abbr} clinches #8 Seed"
            print(f"  {result_8}")
            self.box_scores.append(f"--- {conf} Play-In: 8th Seed — {l_78.abbr} vs {w_910.abbr} ---")
            self.box_scores.extend(game_8.format_box_score())
            self.box_scores.append(result_8)
            self.box_scores.append("")

            playoff_teams[conf] = {
                1: s[0], 2: s[1], 3: s[2], 4: s[3], 5: s[4], 6: s[5], 7: w_78, 8: w_8
            }

        print("\n=== NBA PLAYOFFS ===")
        next_round = {"East": {}, "West": {}}

        print("\n--- FIRST ROUND ---")
        for conf in ["East", "West"]:
            p = playoff_teams[conf]
            matchups = [(1,8), (4,5), (3,6), (2,7)]
            for high, low in matchups:
                winner, hw, lw = self.play_series(p[high], p[low])
                print(f"\n#{high} {p[high].abbr} vs #{low} {p[low].abbr}: {winner.abbr} wins series {max(hw, lw)}-{min(hw, lw)}")
                next_round[conf][high] = winner

        print("\n--- CONFERENCE SEMIFINALS ---")
        conf_finals = {"East": {}, "West": {}}
        for conf in ["East", "West"]:
            n = next_round[conf]
            m1, m1h, m1l = self.play_series(n[1], n[4])
            m2, m2h, m2l = self.play_series(n[2], n[3])
            print(f"\n{n[1].abbr} vs {n[4].abbr}: {m1.abbr} wins series {max(m1h, m1l)}-{min(m1h, m1l)}")
            print(f"{n[2].abbr} vs {n[3].abbr}: {m2.abbr} wins series {max(m2h, m2l)}-{min(m2h, m2l)}")
            conf_finals[conf] = (m1, m2)

        print("\n--- CONFERENCE FINALS ---")
        finals = []
        for conf in ["East", "West"]:
            t1, t2 = conf_finals[conf]
            winner, hw, lw = self.play_series(t1, t2)
            print(f"\n{t1.abbr} vs {t2.abbr}: {winner.abbr} wins series {max(hw, lw)}-{min(hw, lw)} and advances to the Finals!")
            finals.append(winner)

        print("\n==============================")
        print("         NBA FINALS           ")
        print("==============================\n")

        t1, t2 = finals[0], finals[1]
        winner, hw, lw = self.play_series(t1, t2)

        print(f"\n\nCHAMPION: {winner.abbr} WINS THE NBA CHAMPIONSHIP!")

    def display_total_wins(self):
        print("\n==============================")
        print("    TOTAL WINS (ALL GAMES)    ")
        print("==============================\n")
        sorted_teams = sorted(self.teams.values(), key=lambda t: t.abbr)
        for t in sorted_teams:
            print(f"{t.abbr:<5} | {t.wins} Wins")

def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    json_file = os.path.join(script_dir, 'nba_rosters26_v3.json')

    if not os.path.exists(json_file):
        print(f"Error: Could not find nba_rosters26_v3.json at {json_file}")
        return

    with open(json_file, 'r') as f:
        teams_data = json.load(f)

    if len(teams_data) < 30:
        print(f"Error: Your JSON file only has {len(teams_data)} teams. You need all 30 teams mapped to run Season Mode!")
        return

    season = SeasonMode(teams_data)
    season.build_schedule()
    season.sim_regular_season()
    season.display_standings()
    season.display_league_leaders()
    season.run_postseason()
    season.display_total_wins()

    # Write only box scores to Word document
    print("\nWriting box scores to Word document...")
    doc = Document()
    doc.add_heading("N3BL NBA 2026 Season — Box Scores", 0)

    for line in season.box_scores:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    output_path = os.path.join(script_dir, "n3bl_nba26_season_box_scores.docx")
    doc.save(output_path)
    print(f"Word document saved to: {output_path}")

if __name__ == "__main__":
    main()
